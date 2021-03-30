from docx import Document

def add_str_docx(str1, dir_doc):
    doc = Document()
    doc.add_paragraph(str1)
    doc.save(dir_doc)
    return doc

def replace_str_docx(str2):
    add_doc = add_str_docx(str1, dir_doc)
    for paragraph in add_doc.paragraphs:
        if paragraph.text:
            if str1 == paragraph.text:
                paragraph.text = str2
                add_doc.save(dir_doc)
                return 1
            return 0
        continue

dir_doc = '1.docx'
str1 = 'A plain paragraph one!'
str2 = 'A plain paragraph two!'
replace = replace_str_docx(str2)
#add = add_str_docx(str1, dir_doc)